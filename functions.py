import os
from PyPDF4 import PdfFileMerger
from docx import Document

def extractFiles(directory):
    
    """
    It takes directory as input.
    It returns list of all files present in the specified directory.
    """

    try:
        allFiles = []
        for rootDirectory, folder, files in os.walk(directory):
            for file in files:
                allFiles.append(rootDirectory  + "\\" + file)
        return allFiles

    except Exception as e:
        return e


def mergePDF(directory, mergedFilename):
    
    """
    It takes directory, mergedFilename as arguments.
    NOTE: mergedFilename is the name that you want to give to your new file that you are going to create.
    It returns a PDF file after merging all the PDF files found in the specified directory.
    """
    
    try:
        merger = PdfFileMerger()
        pdfFiles = []
        for file in extractFiles(directory):
            if file.endswith('.pdf'):
                pdfFiles.append(file)
        if len(pdfFiles) == 0:
            return 'The directory specified by contains no PDF file!'
        elif len(pdfFiles) == 1:
            return 'The directory specified by you contains single PDF file, hence merging is not possible'
        else:
            for file in pdfFiles:
                merger.append(file)
            merger.write(f"{mergedFilename}.pdf")
            return f'{len(pdfFiles)} PDF files have been merged!'

    except Exception as e:
        return e


def mergeDOCX(directory, mergedFilename):
    
    """
    It takes directory, mergedFilename as arguments.
    NOTE: mergedFilename is the name that you want to give to your new file that you are going to create.
    It returns a doc file after merging all the DOCX files found in the specified directory.
    """
    
    try:
        merger = Document()
        docFiles = []
        for file in extractFiles(directory):
            if file.endswith('.docx'):
                docFiles.append(file)
        if len(docFiles) == 0:
            return 'The directory specified by contains no DOCX file!'
        elif len(docFiles) == 1:
            return 'The directory specified by contains single DOCX file, hence merging is not possible!'
        else:
            for index, file in enumerate(docFiles):
                doc = Document(file)
                for paragraph in doc.paragraphs:
                    merger.add_paragraph(paragraph.text)
                if index < len(docFiles) - 1:
                    merger.add_page_break()
            merger.save(f'{mergedFilename}.docx')
            return f"{len(docFiles)} DOCX files have been merged!"

    except Exception as e:
        return e


def mergeTXT(directory, mergedFilename):
    
    """
    It takes directory, mergedFilename as arguments.
    NOTE: mergedFilename is the name that you want to give to your new file that you are going to create.
    It returns a TXT file after merging all the TXT files found in the specified directory.
    """
    
    try:
        txtFiles = []
        for file in extractFiles(directory):
            if file.endswith('.txt'):
                txtFiles.append(file)
        if len(txtFiles) == 0:
            return 'The directory specified by contains no TXT file!'
        elif len(txtFiles) == 1:
            return 'The directory specified by contains single document file, hence merging is not possible!'
        else:
            with open(f'{mergedFilename}.txt', 'w') as outputFile:
                for file in txtFiles:
                    with open(file, 'r') as f:
                        outputFile.write(f.read() + '\n')
            return f'{len(txtFiles)} TXT files have been merged!'

    except Exception as e:
        return e